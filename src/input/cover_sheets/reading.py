"""
Functions realted to scraping the data from an incoming cover sheet.
"""

import src.utility as utility
from src.constants import COVER_SHEET_DIRECTORY, HEADER_MAP

from docx import Document
from docx.text.paragraph import Paragraph
import pandas as pd
import os

def read_cover_sheet(document):
    """
    Returns a dictionary containing the data collected from the passed word document
    
    :param document: A Document object holding a performance cover sheet from a .docx file
    :return: A dictionary containing all of the relevant data scraped from the passed cover sheet document
    """
    data = {}
    header = None

    for block in utility.iter_block_items(document):
        if isinstance(block, Paragraph):
            if block.text in HEADER_MAP.keys():     # if the paragraph holds one of the headers used to indicate that text can be inputted from user
                header = block.text
        else:   # if block is Table object
            if header:  # if table comes directly after a header
                text_input = None
                for row in block.rows: 
                    for cell in row.cells:
                        for paragraph in cell.paragraphs: 
                            if text_input:  # allows for multiple paragraphs in input
                                text_input = "{} {}".format(text_input, paragraph.text)
                            else:
                                text_input = paragraph.text
                data[HEADER_MAP[header]] = text_input
                header = None
            else:
                for row in block.rows:
                    row_text = []
                    for cell in row.cells:
                        for paragraph in cell.paragraphs: 
                            row_text.append(paragraph.text)     # collects all paragraphs in a given cell into a list

                    # storing of checkboxes in data field - continues until all checkboxes are retrieved from cell
                    while any([i in ["☒", "☐"] for i in row_text]):
                        try:
                            checked_index = row_text.index("☒")
                        except:
                            checked_index = -1

                        try:
                            unchecked_index = row_text.index("☐")
                        except:
                            unchecked_index = -1

                        index = max(checked_index, unchecked_index)
                        checkbox = row_text.pop(index)
                        column_title = row_text.pop(index)

                        data[column_title] = 1 if checkbox == "☒" else 0   # stores a 1 if box is checked, else 0
    
    return data

def process_cover_sheets(cover_sheets_list):
    """
    Creates a DataFrame object from a list of cover sheets
    
    :param cover_sheets_list: A list of Document objects, each of which was created from a separate .docx file containing a cover sheet
    :return: A DataFrame object with each row representing a cover sheet.
    """
    data = []
    for cover_sheet in cover_sheets_list:
        data.append(read_cover_sheet(cover_sheet))
        
    return pd.DataFrame(data)

def get_cover_sheets(path=None):
    """
    Returns a list of docx Document objects representing cover sheets, each of which are retrieved from the folder where cover sheets are stored.

    :param path: The path to the directory where cover sheet objects are stored. NOTE: This directory should only include cover sheets.
    :return: A list of docx Document objects representing cover sheets, each of which are retrieved from the folder where cover sheets are stored.
    """
    cover_sheets = []

    if path == None:
        path = COVER_SHEET_DIRECTORY    # retrieves cover sheets from the default directory specified by a constant if no path was passed

    try:
        for filename in os.listdir(path):  # retrieves all file names from cover sheet directory
            cover_sheets.append(Document(f"{path}{filename}"))     # creates Document objects for every cover sheet
    except FileNotFoundError as e:
        print(f"Unable to retrieve cover sheets from path {path}")
        return None

    return cover_sheets